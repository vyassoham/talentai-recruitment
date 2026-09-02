import os
from typing import Optional
from core.storage import LocalStorage

class ExtractionError(Exception):
    pass

class DocumentExtractor:
    def __init__(self, storage: LocalStorage):
        self.storage = storage

    def extract_text(self, storage_key: str) -> str:
        """
        Extracts raw text from the stored document.
        Supports PDF and DOCX.
        """
        file_path = self.storage.get(storage_key)
        if not os.path.exists(file_path):
            raise ExtractionError(f"File not found: {storage_key}")

        ext = os.path.splitext(storage_key)[1].lower()
        
        try:
            if ext == '.pdf':
                return self._extract_pdf(file_path)
            elif ext == '.docx':
                return self._extract_docx(file_path)
            elif ext in ['.txt', '.md', '.csv']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            else:
                raise ExtractionError(f"Unsupported extraction format: {ext}")
        except Exception as e:
            raise ExtractionError(f"Failed to extract text: {str(e)}")

    def _extract_pdf(self, file_path: str) -> str:
        import fitz # PyMuPDF
        text_parts = []
        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    text_parts.append(page.get_text())
            if not text_parts:
                raise ExtractionError("PDF appears to be empty or scanned without OCR.")
            return "\n\n".join(text_parts)
        except Exception as e:
            raise ExtractionError(f"PyMuPDF Error: {e}")

    def _extract_docx(self, file_path: str) -> str:
        import docx
        text_parts = []
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            return "\n".join(text_parts)
        except Exception as e:
            raise ExtractionError(f"python-docx Error: {e}")
